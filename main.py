import easyocr
import cv2
import numpy as np
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX  # Correct import for WD_COLOR_INDEX
from spellchecker import SpellChecker

def detect_highlights(image):
    """
    Detect highlighted areas in the image (yellow and green highlights).
    """
    hsv = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)  # Convert image to HSV
    
    # Define the yellow and green color ranges in HSV
    lower_yellow = np.array([20, 100, 100])
    upper_yellow = np.array([30, 255, 255])
    lower_green = np.array([40, 40, 40])
    upper_green = np.array([90, 255, 255])
    
    # Create masks for yellow and green regions (highlighted areas)
    mask_yellow = cv2.inRange(hsv, lower_yellow, upper_yellow)
    mask_green = cv2.inRange(hsv, lower_green, upper_green)

    # Find contours for yellow and green highlighted areas
    contours_yellow, _ = cv2.findContours(mask_yellow, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    contours_green, _ = cv2.findContours(mask_green, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    
    highlight_boxes = []
    
    # Process yellow highlights
    for contour in contours_yellow:
        if cv2.contourArea(contour) > 100:  # Filter small noise
            x, y, w, h = cv2.boundingRect(contour)
            highlight_boxes.append((x, y, x + w, y + h, 'yellow'))  # (x1, y1, x2, y2, color)

    # Process green highlights
    for contour in contours_green:
        if cv2.contourArea(contour) > 100:  # Filter small noise
            x, y, w, h = cv2.boundingRect(contour)
            highlight_boxes.append((x, y, x + w, y + h, 'green'))  # (x1, y1, x2, y2, color)
    
    return highlight_boxes

def extract_text_with_easyocr(image_path):
    """
    Use EasyOCR to extract text and bounding boxes.
    """
    reader = easyocr.Reader(['en'])
    result = reader.readtext(image_path)
    
    ocr_data = []
    for detection in result:
        text = detection[1]
        bbox = detection[0]  # Bounding box for the detected text
        
        spell = SpellChecker()
        text = spell.correction(text)
        ocr_data.append({"text": text, "bbox": bbox})
    
    return ocr_data

def create_docx_with_highlighted_text(ocr_data, highlight_boxes, output_path):
    """
    Create a DOCX document with highlighted text based on bounding boxes.
    Differentiate yellow and green highlights.
    """
    doc = Document()
    
    # Iterate over the OCR data and check if text is in a highlighted region
    for data in ocr_data:
        text = data['text']
        bbox = data['bbox']
        
        # Check if any highlight region overlaps with the OCR bounding box
        is_highlighted = False
        highlight_color = None
        
        for hbox in highlight_boxes:
            x1, y1, x2, y2, color = hbox
            # Check if the bounding boxes overlap
            if not (bbox[2][0] < x1 or bbox[0][0] > x2 or bbox[2][1] < y1 or bbox[0][1] > y2):
                is_highlighted = True
                highlight_color = color
                break
        
        # Add the text to DOCX and highlight if necessary
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text + " ")
        
        # Highlight the text based on the color
        if is_highlighted:
            if highlight_color == 'yellow':
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # Yellow highlight
            elif highlight_color == 'green':
                run.font.highlight_color = 4  # Green highlight

    doc.save(output_path)
    print(f"Document saved as {output_path}")

def main(image_path, output_path):
    # Load the image with OpenCV
    image = cv2.imread(image_path)
    
    # Detect highlighted regions in the image
    highlight_boxes = detect_highlights(image)
    
    # Extract text and bounding boxes using EasyOCR
    ocr_data = extract_text_with_easyocr(image_path)
    
    # Create DOCX with highlighted text
    create_docx_with_highlighted_text(ocr_data, highlight_boxes, output_path)

# Input: Image path and Output DOCX path
image_path = 'sample10f.jpg'  # Replace with your image path
output_path = 'output.docx'  # Desired output DOCX file name

# Run the process
main(image_path, output_path)
